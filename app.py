import os
import re
import ast
import io
import joblib
import pandas as pd
import streamlit as st

try:
    from groq import Groq
except ImportError:
    Groq = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    SimpleDocTemplate = None

# -------------------------
# PAGE CONFIG
# -------------------------
st.set_page_config(
    page_title="AI Resume Matcher",
    page_icon="💼",
    layout="wide"
)

# -------------------------
# COLORS
# -------------------------
PRIMARY = "#FF2D2D"
SECONDARY = "#B30000"
BG = "#0E0E0E"
CARD = "#1A1A1A"
TEXT = "#FFFFFF"
MUTED = "#B3B3B3"
BORDER = "#2A2A2A"

# -------------------------
# CSS
# -------------------------
st.markdown(
    f"""
    <style>
    .stApp {{
        background-color: {BG};
        color: {TEXT};
    }}

    .block-container {{
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }}

    h1, h2, h3, h4 {{
        color: {PRIMARY};
    }}

    p, label, div, span {{
        color: {TEXT};
    }}

    .hero-card {{
        background: linear-gradient(135deg, #151515 0%, #1d0f0f 100%);
        border: 1px solid {BORDER};
        border-radius: 22px;
        padding: 1.6rem;
        margin-bottom: 1.2rem;
    }}

    .info-card, .result-card {{
        background: {CARD};
        border: 1px solid {BORDER};
        border-radius: 18px;
        padding: 1rem 1.15rem;
        margin-top: 1rem;
    }}

    div[data-testid="stFileUploader"] {{
        background: {CARD};
        border: 1px solid {BORDER};
        border-radius: 16px;
        padding: 0.75rem;
    }}

    div[data-testid="stMetric"] {{
        background: {CARD};
        border: 1px solid {BORDER};
        border-radius: 16px;
        padding: 0.75rem;
    }}

    .stButton > button {{
        background: linear-gradient(90deg, {PRIMARY} 0%, {SECONDARY} 100%);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 0.7rem 1.1rem;
        font-weight: 700;
        width: 100%;
    }}

    .stButton > button:hover {{
        filter: brightness(1.08);
        color: white;
    }}

    div[data-testid="stProgressBar"] > div > div > div > div {{
        background: linear-gradient(90deg, {PRIMARY} 0%, {SECONDARY} 100%);
    }}

    .section-label {{
        font-size: 0.95rem;
        font-weight: 700;
        color: {PRIMARY};
        margin-bottom: 0.4rem;
        text-transform: uppercase;
        letter-spacing: 0.04em;
    }}

    .muted {{
        color: {MUTED};
        font-size: 0.95rem;
    }}

    .pill {{
        display: inline-block;
        padding: 0.35rem 0.7rem;
        margin: 0.2rem 0.35rem 0.2rem 0;
        border-radius: 999px;
        font-size: 0.88rem;
        font-weight: 600;
        border: 1px solid {BORDER};
        background: #121212;
        color: white;
    }}

    .pill-match {{
        background: rgba(29, 185, 84, 0.12);
        border: 1px solid rgba(29, 185, 84, 0.35);
    }}

    .pill-missing {{
        background: rgba(255, 45, 45, 0.12);
        border: 1px solid rgba(255, 45, 45, 0.35);
    }}

    .small-note {{
        font-size: 0.85rem;
        color: {MUTED};
    }}
    </style>
    """,
    unsafe_allow_html=True
)

# -------------------------
# PATHS
# -------------------------
MODEL_PATH = "resume_job_matcher_model.pkl"
TRAINING_PATH = "training_pairs_debug.csv"
SKILLS_PATH = "skills.txt"

# -------------------------
# LOAD MODEL/DATA
# -------------------------
@st.cache_resource
def load_model():
    return joblib.load(MODEL_PATH)

@st.cache_data
def load_training_data():
    return pd.read_csv(TRAINING_PATH)

@st.cache_data
def load_skills():
    try:
        with open(SKILLS_PATH, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read().lower()
        skills = re.split(r"[\n,]+", raw)
        return sorted(list(set([s.strip() for s in skills if s.strip()])))
    except FileNotFoundError:
        return []

model = load_model()
training_df = load_training_data()

# -------------------------
# GROQ
# -------------------------
def get_groq_client():
    api_key = None

    if "GROQ_API_KEY" in st.secrets:
        api_key = st.secrets["GROQ_API_KEY"]
    elif os.getenv("GROQ_API_KEY"):
        api_key = os.getenv("GROQ_API_KEY")

    if not api_key or Groq is None:
        return None

    return Groq(api_key=api_key)

# -------------------------
# HELPERS
# -------------------------
def clean_text(text):
    text = str(text).lower()
    text = re.sub(r"[^a-zA-Z0-9\s\+\#\.\-/]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def predict_match(resume_text, job_text):
    combined = "resume: " + clean_text(resume_text) + " job: " + clean_text(job_text)
    pred = model.predict([combined])[0]
    prob = model.predict_proba([combined])[0][1]
    return int(pred), float(prob)

def extract_skills_from_text(text, skills_list):
    cleaned = f" {clean_text(text)} "
    found = set()

    for skill in skills_list:
        skill_clean = skill.strip().lower()
        if not skill_clean:
            continue

        pattern = r"(?<!\w)" + re.escape(skill_clean) + r"(?!\w)"
        if re.search(pattern, cleaned):
            found.add(skill_clean)

    return sorted(found)

def compare_uploaded_skills(resume_text, job_text):
    skills_list = load_skills()
    resume_skills = extract_skills_from_text(resume_text, skills_list)
    job_skills = extract_skills_from_text(job_text, skills_list)
    matched_skills = sorted(set(resume_skills).intersection(set(job_skills)))
    missing_skills = sorted(set(job_skills) - set(resume_skills))
    return resume_skills, job_skills, matched_skills, missing_skills

def decode_text_bytes(raw):
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")

def read_txt_file(uploaded_file):
    return decode_text_bytes(uploaded_file.read())

def read_pdf_file(uploaded_file):
    if PyPDF2 is None:
        raise ImportError("PyPDF2 is not installed. Add PyPDF2 to requirements.txt")

    reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
    text = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text.append(page_text)

    return "\n".join(text).strip()

def read_docx_file(uploaded_file):
    if Document is None:
        raise ImportError("python-docx is not installed. Add python-docx to requirements.txt")

    doc = Document(io.BytesIO(uploaded_file.read()))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()

def read_uploaded_file(uploaded_file):
    name = uploaded_file.name.lower()

    if name.endswith(".txt"):
        return read_txt_file(uploaded_file)
    if name.endswith(".pdf"):
        return read_pdf_file(uploaded_file)
    if name.endswith(".docx"):
        return read_docx_file(uploaded_file)

    raise ValueError("Unsupported file type. Upload TXT, PDF, or DOCX.")

def render_skill_pills(skills, pill_class):
    if not skills:
        st.write("None found")
        return

    html = "".join([f'<span class="pill {pill_class}">{skill}</span>' for skill in skills])
    st.markdown(html, unsafe_allow_html=True)

def get_match_label(score):
    if score >= 0.75:
        return "Strong Match"
    if score >= 0.45:
        return "Moderate Match"
    return "Low Match"

def build_resume_rewrite_prompt(resume_text, job_text, score, matched_skills, missing_skills):
    matched = ", ".join(matched_skills[:20]) if matched_skills else "None identified"
    missing = ", ".join(missing_skills[:20]) if missing_skills else "None identified"

    return f"""
You are an expert resume writer and career coach.

Only use information already present in the resume.
Do not invent employers, job titles, tools, metrics, dates, achievements, or experiences.

Match Score: {round(score * 100, 2)}%
Matched Skills: {matched}
Missing Skills: {missing}

Resume:
{resume_text[:5000]}

Job Description:
{job_text[:5000]}

Return exactly these sections:

## Overall Fit
Write 2 concise sentences.

## Improved Resume Lines
Write 5 improved bullet points tailored to the job.
Each bullet must start with a strong action verb and stay truthful to the resume.

## Keywords To Add If Truthful
List 8 keywords from the job description.

## Top Gaps
List the top 3 gaps between the resume and the role.
"""

def generate_ai_feedback(resume_text, job_text, score, matched_skills, missing_skills):
    client = get_groq_client()

    if client is None:
        return None

    prompt = build_resume_rewrite_prompt(
        resume_text,
        job_text,
        score,
        matched_skills,
        missing_skills
    )

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "You are a precise resume rewriting assistant. You never invent experience."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3
        )

        return response.choices[0].message.content

    except Exception as e:
        error_text = str(e)
        if "invalid_api_key" in error_text.lower() or "invalid api key" in error_text.lower():
            return "Groq rejected the API key. Update your GROQ_API_KEY in Streamlit secrets and redeploy."
        return f"AI feedback could not be generated: {e}"

def create_docx_download(ai_text, score, matched_skills, missing_skills):
    doc = Document()
    doc.add_heading("AI Resume Match Feedback", level=1)

    doc.add_paragraph(f"Match Score: {round(score * 100, 2)}%")
    doc.add_paragraph(f"Matched Skills: {', '.join(matched_skills) if matched_skills else 'None found'}")
    doc.add_paragraph(f"Missing Skills: {', '.join(missing_skills) if missing_skills else 'None found'}")

    doc.add_heading("AI Resume Rewrite Suggestions", level=2)

    for line in ai_text.split("\n"):
        clean_line = line.strip()

        if not clean_line:
            continue

        if clean_line.startswith("## "):
            doc.add_heading(clean_line.replace("## ", ""), level=2)
        elif clean_line.startswith("- "):
            doc.add_paragraph(clean_line.replace("- ", ""), style="List Bullet")
        else:
            doc.add_paragraph(clean_line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_download(ai_text, score, matched_skills, missing_skills):
    if SimpleDocTemplate is None:
        return None

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("AI Resume Match Feedback", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Match Score: {round(score * 100, 2)}%", styles["Normal"]))
    story.append(Paragraph(f"Matched Skills: {', '.join(matched_skills) if matched_skills else 'None found'}", styles["Normal"]))
    story.append(Paragraph(f"Missing Skills: {', '.join(missing_skills) if missing_skills else 'None found'}", styles["Normal"]))
    story.append(Spacer(1, 12))

    for line in ai_text.split("\n"):
        clean_line = line.strip()

        if not clean_line:
            story.append(Spacer(1, 8))
            continue

        clean_line = clean_line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if clean_line.startswith("## "):
            story.append(Paragraph(clean_line.replace("## ", ""), styles["Heading2"]))
        elif clean_line.startswith("- "):
            story.append(Paragraph("• " + clean_line.replace("- ", ""), styles["Normal"]))
        else:
            story.append(Paragraph(clean_line, styles["Normal"]))

    pdf.build(story)
    buffer.seek(0)
    return buffer

# -------------------------
# HEADER
# -------------------------
st.markdown(
    """
    <div class="hero-card">
        <h1 style="margin-bottom:0.4rem;">💼 AI Resume Matcher</h1>
        <p style="margin:0; font-size:1.05rem;">Match smarter. Apply better.</p>
        <p class="muted" style="margin-top:0.6rem;">
            Upload a resume and a job description to get a match score, personalized skill insights, AI-tailored resume bullet rewrites, and downloadable feedback.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

# -------------------------
# UPLOAD SECTION
# -------------------------
left, right = st.columns(2)

with left:
    st.markdown('<div class="section-label">Resume Upload</div>', unsafe_allow_html=True)
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["txt", "pdf", "docx"],
        label_visibility="collapsed"
    )

with right:
    st.markdown('<div class="section-label">Job Description Upload</div>', unsafe_allow_html=True)
    job_file = st.file_uploader(
        "Upload Job Description",
        type=["txt", "pdf", "docx"],
        label_visibility="collapsed"
    )

# -------------------------
# MAIN APP
# -------------------------
if resume_file and job_file:
    try:
        resume_text = read_uploaded_file(resume_file)
        job_text = read_uploaded_file(job_file)

        if not resume_text.strip():
            st.error("The uploaded resume appears empty or could not be read.")
            st.stop()

        if not job_text.strip():
            st.error("The uploaded job description appears empty or could not be read.")
            st.stop()

    except Exception as e:
        st.error(f"File reading error: {e}")
        st.stop()

    pred, score = predict_match(resume_text, job_text)
    label = get_match_label(score)

    resume_skills, job_skills, matched_skills, missing_skills = compare_uploaded_skills(
        resume_text,
        job_text
    )

    st.markdown('<div class="section-label">Match Results</div>', unsafe_allow_html=True)

    st.progress(min(max(score, 0.0), 1.0))
    st.caption(f"Overall match confidence: {round(score * 100, 2)}%")

    c1, c2, c3 = st.columns(3)

    with c1:
        st.metric("Match Score", f"{round(score * 100, 2)}%")

    with c2:
        st.metric("Prediction", label)

    with c3:
        st.metric("Matched Skills", f"{len(matched_skills)} / {len(job_skills)}")

    skill_col1, skill_col2 = st.columns(2)

    with skill_col1:
        st.markdown('<div class="result-card">', unsafe_allow_html=True)
        st.subheader("✅ Matched Skills")
        render_skill_pills(matched_skills[:25], "pill-match")
        st.markdown("</div>", unsafe_allow_html=True)

    with skill_col2:
        st.markdown('<div class="result-card">', unsafe_allow_html=True)
        st.subheader("❌ Missing Skills")
        render_skill_pills(missing_skills[:25], "pill-missing")
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="result-card">', unsafe_allow_html=True)
    st.subheader("🧠 Quick Feedback")

    if score > 0.75:
        st.success("Strong alignment. This resume appears to fit the role well.")
    elif score > 0.45:
        st.warning("Decent alignment, but there is room to tailor the resume more directly to the job.")
    else:
        st.error("Lower alignment. The resume likely needs stronger tailoring for this position.")

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-label">AI Resume Rewrite</div>', unsafe_allow_html=True)

    ai_col1, ai_col2 = st.columns([2, 1])

    with ai_col1:
        generate_clicked = st.button("Generate AI Rewrite Suggestions")

    with ai_col2:
        groq_ready = get_groq_client() is not None
        st.markdown(
            f"<p class='small-note'>{'Groq connected' if groq_ready else 'Groq key not detected'}</p>",
            unsafe_allow_html=True
        )

    if "ai_output" not in st.session_state:
        st.session_state.ai_output = None

    if generate_clicked:
        with st.spinner("Generating tailored resume improvements..."):
            st.session_state.ai_output = generate_ai_feedback(
                resume_text,
                job_text,
                score,
                matched_skills,
                missing_skills
            )

    if st.session_state.ai_output:
        ai_output = st.session_state.ai_output

        st.markdown('<div class="result-card">', unsafe_allow_html=True)
        st.markdown(ai_output)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="section-label">Download Feedback</div>', unsafe_allow_html=True)

        docx_file = create_docx_download(
            ai_output,
            score,
            matched_skills,
            missing_skills
        )

        download_col1, download_col2 = st.columns(2)

        with download_col1:
            st.download_button(
                label="⬇️ Download Word Feedback",
                data=docx_file,
                file_name="ai_resume_feedback.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with download_col2:
            pdf_file = create_pdf_download(
                ai_output,
                score,
                matched_skills,
                missing_skills
            )

            if pdf_file:
                st.download_button(
                    label="⬇️ Download PDF Feedback",
                    data=pdf_file,
                    file_name="ai_resume_feedback.pdf",
                    mime="application/pdf"
                )
            else:
                st.info("PDF download requires reportlab in requirements.txt.")

    elif generate_clicked and not st.session_state.ai_output:
        st.info("Add a Groq API key in Streamlit secrets to enable AI-generated rewrites.")

    with st.expander("📄 View Uploaded Resume Text"):
        st.write(resume_text[:4000])

    with st.expander("📄 View Uploaded Job Description"):
        st.write(job_text[:4000])

else:
    st.markdown(
        """
        <div class="info-card">
            <h3 style="margin-top:0;">How to use</h3>
            <p>1. Upload your resume as TXT, PDF, or DOCX.</p>
            <p>2. Upload the job description as TXT, PDF, or DOCX.</p>
            <p>3. Review your match score and personalized skill analysis.</p>
            <p>4. Generate AI rewrite suggestions for stronger resume bullets.</p>
            <p>5. Download your feedback as a Word document or PDF.</p>
        </div>
        """,
        unsafe_allow_html=True
    )
